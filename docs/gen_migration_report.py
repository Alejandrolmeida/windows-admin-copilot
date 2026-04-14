#!/usr/bin/env python3
"""
gen_migration_report.py
Genera el informe Word de planificacion de migracion Azure para Plenergy.

Lee datos de monitorización desde docs/data/ (descargados por Collect-MonitoringData.ps1).
No contiene credenciales ni datos de servidor hardcodeados.

Requiere:
    pip install python-docx

Uso:
    python docs/gen_migration_report.py
    python docs/gen_migration_report.py --output "docs/MiInforme.docx"
    python docs/gen_migration_report.py --data-dir docs/data --cycle checkpoint
"""

import sys
import json
import argparse
from pathlib import Path
from datetime import datetime, date
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx no instalado.")
    print("Ejecuta: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Rutas (relativas al repo root)
# ---------------------------------------------------------------------------
SCRIPT_DIR  = Path(__file__).parent                    # docs/
REPO_ROOT   = SCRIPT_DIR.parent                       # raiz del repo
CONFIG_DIR  = REPO_ROOT / '.config'                   # gitignored
DEFAULT_DATA_DIR = SCRIPT_DIR / 'data'                # gitignored

# ---------------------------------------------------------------------------
# Conocimiento de los informes previos (13/04/2026)
# Hardcodeado porque los .docx son conocidos y están gitignored
# ---------------------------------------------------------------------------
PREV_REPORTS = [
    {
        "title": "Fragmentacion critica en srvplenoilfs (13/04/2026)",
        "file":  "Informe-FragmentacionSQL-srvplenoilfs-20260413.docx",
        "findings": [
            "Causa raiz identificada: fragmentacion del 98,46% en el indice PK_HechosVentasPOS.",
            "Tabla afectada: HechosVentasPOS — 147 millones de filas, 57 GB de datos.",
            "Consecuencia directa: el job ProcesarCubosVentasNoche tarda >10 horas en lugar de las 2-3 horas esperadas.",
            "No existen jobs de mantenimiento de indices configurados en SQL Agent.",
            "Recomendacion inmediata: REBUILD INDEX con ONLINE=ON en ventana de baja actividad.",
        ],
        "risk": "CRITICO",
    },
    {
        "title": "Incidente DBCC SHRINKDATABASE — Gobernanza (13/04/2026)",
        "file":  "Informe-DBCC-SHRINKDATABASE-srvplenoilfs-20260413.docx",
        "findings": [
            "El usuario rafael.lazaga ejecuto DBCC SHRINKDATABASE en la base de datos de produccion sin autorizacion.",
            "El SHRINK agrava la fragmentacion ya existente, creando un ciclo de degradacion recurrente.",
            "El benchmark de SQL Server fue inicialmente sospechoso pero quedo DESCARTADO como causa recurrente.",
            "Se confirma que la lentitud del job nocturno ocurrio en noches sin benchmark activo.",
            "El SHRINK se ejecuta de forma recurrente por un ciclo de automatizacion no autorizado.",
        ],
        "risk": "ALTO",
    },
    {
        "title": "Auditoria de gobernanza y seguridad SQL (13/04/2026)",
        "file":  "Informe-Gobernanza-srvplenoilfs-20260413.docx",
        "findings": [
            "16 logins con rol sysadmin — excesivo para las mejores practicas de minimo privilegio.",
            "Login 'sa' habilitado; login 'prueba' habilitado con permisos amplios.",
            "rafael.lazaga anadido como sysadmin el 23/03/2026 sin registro de cambio documentado.",
            "0 alertas de SQL Agent configuradas para eventos criticos (errores 823, 824, 825, 17806).",
            "Ciclo SHRINK nocturno ejecutado por automatizacion no documentada.",
            "Recomendacion: implementar control de cambios, deshabilitar sa, reducir sysadmins.",
        ],
        "risk": "ALTO",
    },
]

# ---------------------------------------------------------------------------
# Jobs SQL Agent criticos identificados
# ---------------------------------------------------------------------------
SQL_JOBS = [
    {"name": "ProcesarCubosVentasNoche",   "server": "srvplenoilfs",   "schedule": "Diario 01:30",        "duration_h": 10.6, "captured": True,  "priority": "CRITICO"},
    {"name": "ProcesarCubosContabilidad",  "server": "srvplenoilfs",   "schedule": "Diario 22:45",        "duration_h": 1.3,  "captured": True,  "priority": "MEDIO"},
    {"name": "ProcesaCubosCompra",         "server": "srvplenoilfs",   "schedule": "Diario 09:00/17:00",  "duration_h": 2.1,  "captured": True,  "priority": "MEDIO"},
    {"name": "ProcesarCubosdomingo",       "server": "srvplenoilfs",   "schedule": "Sab+Lun 01:30",       "duration_h": 6.4,  "captured": False, "priority": "ALTO"},
    {"name": "DatabaseBackup FULL",        "server": "sql-restore-001","schedule": "Diario 00:00",        "duration_h": 3.3,  "captured": True,  "priority": "ALTO"},
    {"name": "IndicesEstadisticas",        "server": "sql-restore-001","schedule": "Dom/Mar/Mie/Vie 18:30","duration_h": 25.7, "captured": False, "priority": "CRITICO"},
]

# ---------------------------------------------------------------------------
# Helpers de formato Word
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    """Establece color de fondo en una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, style='List Bullet'):
    p = doc.add_paragraph(text, style=style)
    return p


def add_table_row(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val) if val is not None else ''
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        if bg:
            set_cell_bg(cell, bg)
    return row


# ---------------------------------------------------------------------------
# Analisis de datos de monitorización
# ---------------------------------------------------------------------------
def load_monitoring_data(data_dir: Path, server_filter: str = None):
    """Carga y parsea los JSON de monitorización de docs/data/."""
    results = {}

    if not data_dir.exists():
        print(f"WARN: Directorio de datos no encontrado: {data_dir}")
        return results

    for json_file in sorted(data_dir.glob('*.json')):
        if json_file.name.startswith('.'):
            continue
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)

            server = data.get('server', json_file.stem.split('_')[0])
            if server_filter and server_filter.lower() not in server.lower():
                continue

            results[server] = {
                'file':      json_file.name,
                'raw':       data,
                'stats':     compute_stats(data),
            }
            print(f"  [OK] Cargado: {json_file.name} — server={server}, samples={data.get('samples_collected', '?')}")
        except Exception as e:
            print(f"  [WARN] No se pudo parsear {json_file.name}: {e}")

    return results


def compute_stats(data: dict) -> dict:
    """Calcula estadisticas agregadas de los samples de monitorización."""
    samples = data.get('samples', [])
    if not samples:
        return {}

    # CPU
    sql_cpu  = [s.get('cpu', {}).get('sql_cpu_utilization', 0) for s in samples if s.get('cpu')]
    tot_cpu  = [s.get('cpu', {}).get('sql_cpu_utilization', 0) + s.get('cpu', {}).get('other_cpu', 0)
                for s in samples if s.get('cpu')]

    # Memoria (MB)
    buf_pool_mb  = [s.get('memory', {}).get('buffer_pool_mb', 0) for s in samples if s.get('memory')]
    total_mem_mb = [s.get('memory', {}).get('total_server_memory_mb', 0) for s in samples if s.get('memory')]
    target_mb    = [s.get('memory', {}).get('target_server_memory_mb', 0) for s in samples if s.get('memory')]
    grants_pend  = [s.get('memory', {}).get('memory_grants_pending', 0) for s in samples if s.get('memory')]

    # Actividad
    connections  = [s.get('activity', {}).get('active_connections', 0) for s in samples if s.get('activity')]
    requests     = [s.get('activity', {}).get('active_requests', 0)    for s in samples if s.get('activity')]
    blocked      = [s.get('activity', {}).get('blocked_processes', 0)  for s in samples if s.get('activity')]

    # Wait stats acumuladas
    wait_totals = defaultdict(float)
    for s in samples:
        for w in s.get('waits', []):
            wait_totals[w['wait_type']] += w.get('wait_time_ms', 0)

    top_waits = sorted(wait_totals.items(), key=lambda x: x[1], reverse=True)[:8]

    def safe_avg(lst): return round(sum(lst) / len(lst), 1) if lst else 0
    def safe_max(lst): return round(max(lst), 1) if lst else 0
    def safe_min(lst): return round(min(lst), 1) if lst else 0

    # Calcular duracion real
    start_time = data.get('start_time', '')
    end_time   = data.get('checkpoint_time', '')
    elapsed_min = data.get('elapsed_minutes', 0)
    if not elapsed_min and samples:
        try:
            t0 = datetime.fromisoformat(samples[0]['timestamp'])
            t1 = datetime.fromisoformat(samples[-1]['timestamp'])
            elapsed_min = round((t1 - t0).total_seconds() / 60, 0)
        except Exception:
            pass

    return {
        'start_time':        start_time,
        'end_time':          end_time,
        'elapsed_min':       elapsed_min,
        'samples_n':         len(samples),
        # CPU
        'cpu_sql_avg':       safe_avg(sql_cpu),
        'cpu_sql_max':       safe_max(sql_cpu),
        'cpu_total_avg':     safe_avg(tot_cpu),
        'cpu_total_max':     safe_max(tot_cpu),
        # Memoria
        'buf_pool_avg_gb':   round(safe_avg(buf_pool_mb)  / 1024, 1),
        'buf_pool_max_gb':   round(safe_max(buf_pool_mb)  / 1024, 1),
        'total_mem_max_gb':  round(safe_max(total_mem_mb) / 1024, 1),
        'target_mem_gb':     round(safe_avg(target_mb)    / 1024, 1),
        'mem_grants_max':    safe_max(grants_pend),
        # Actividad
        'conn_avg':          safe_avg(connections),
        'conn_max':          safe_max(connections),
        'req_avg':           safe_avg(requests),
        'req_max':           safe_max(requests),
        'blocked_max':       safe_max(blocked),
        'blocked_events':    sum(1 for b in blocked if b > 0),
        # Waits
        'top_waits':         top_waits,
    }


# ---------------------------------------------------------------------------
# Generacion del documento Word
# ---------------------------------------------------------------------------
def generate_report(monitoring_data: dict, output_path: Path):
    doc = Document()

    # Estilos globales
    style      = doc.styles['Normal']
    font       = style.font
    font.name  = 'Calibri'
    font.size  = Pt(11)

    # -----------------------------------------------------------------------
    # PORTADA
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Planificacion de Migracion a Azure (IaaS)')
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x47, 0x9D)  # Azul corporativo

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run('Analisis de Workload — SQL Server, Capa de Aplicacion NAV y AD DS\n').font.size = Pt(14)

    client_p = doc.add_paragraph()
    client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = client_p.add_run('Cliente: Plenergy')
    r.bold = True
    r.font.size = Pt(13)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f'Fecha: {date.today().strftime("%d/%m/%Y")}').font.size = Pt(11)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # 1. RESUMEN EJECUTIVO
    # -----------------------------------------------------------------------
    add_heading(doc, '1. Resumen Ejecutivo', 1)

    summary_bullets = [
        'Se han realizado los primeros analisis de infraestructura SQL en los servidores de produccion de Plenergy '
        '(srvplenoilfs y sql-restore-001).',
        'Se identifico la CAUSA RAIZ de la lentitud nocturna recurrente: fragmentacion critica del 98,46% en el '
        'indice principal de HechosVentasPOS (147M filas, 57 GB).',
        'Se completó un primer Ciclo de Monitorización de 24 horas (13-14/04/2026) que captura los principales '
        'jobs diarios de negocio.',
        'Las 24 horas de monitorización NO son suficientemente representativas: quedan sin capturar el job '
        'IndicesEstadisticas (25h 43min, sql-restore-001) y ProcesarCubosdomingo (6h 21min, srvplenoilfs, fines '
        'de semana), ademas de un pico de fin de mes generado por los usuarios.',
        'El ALCANCE DE MIGRACION se ha ampliado para incluir la totalidad de la infraestructura implicada: '
        '2 servidores SQL Server + 2 servidores de aplicacion Dynamics NAV + 1 replica AD DS en Azure. '
        'El modelo adoptado es IaaS (lift & shift) para garantizar compatibilidad total con los procesos '
        'y procedimientos almacenados existentes.',
        'Se propone una monitorización extendida de al menos 7 dias para obtener una fotografia completa del '
        'workload real antes de proceder al sizing definitivo de las instancias Azure IaaS.',
    ]
    for b in summary_bullets:
        add_bullet(doc, b)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Factor de ajuste recomendado para el sizing: ')
    r.bold = True
    p.add_run('x1,3 sobre los picos capturados hasta ahora, hasta disponer de datos de fin de semana y fin de mes.')

    # -----------------------------------------------------------------------
    # 2. CONTEXTO DEL PROYECTO
    # -----------------------------------------------------------------------
    add_heading(doc, '2. Contexto del Proyecto', 1)
    doc.add_paragraph(
        'Plenergy dispone de una infraestructura on-premises compuesta por servidores SQL Server con cargas '
        'criticas de negocio (procesamiento de cubos OLAP, backup diario, actualizacion de indices), '
        'servidores de aplicacion Dynamics NAV y un controlador de dominio Active Directory. '
        'El objetivo es migrar toda esta infraestructura a Azure con el modelo IaaS (Infrastructure as a Service / '
        'lift & shift), manteniendo las VMs en Azure sin cambios en el software ni en la logica de negocio.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Razon de la eleccion de IaaS frente a PaaS: ')
    r.bold = True
    p.add_run(
        'El entorno de Plenergy contiene procedimientos almacenados y procesos de negocio en Dynamics NAV '
        'con dependencias especificas de la version y configuracion actuales. La migracion a PaaS '
        '(Azure SQL Managed Instance, Azure App Service) exigiria un analisis exhaustivo de compatibilidad '
        'y posibles adaptaciones de codigo. El modelo IaaS elimina este riesgo al replicar fielmente el '
        'entorno on-premises en Azure, permitiendo una migracion rapida y segura sin afectar a los procesos.'
    )

    add_heading(doc, '2.1 Entorno actual — Alcance completo de migracion', 2)
    doc.add_paragraph(
        'La siguiente tabla detalla la totalidad de los servidores incluidos en el alcance de migracion a Azure IaaS:'
    )
    servers_tbl = doc.add_table(rows=1, cols=4)
    servers_tbl.style = 'Table Grid'
    hdr = servers_tbl.rows[0].cells
    for i, h in enumerate(['Servidor', 'Rol', 'RAM Total', 'Carga principal']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], 'D6E4F0')

    server_rows = [
        ['srvplenoilfs',    'SQL Server — Procesamiento OLAP / ETL',           '240 GB',       'Cubos de ventas y compras (Navision/BC)'],
        ['sql-restore-001', 'SQL Server — Operacional / Backup-Restore',        '240 GB',       'Base de datos principal + mantenimiento de indices'],
        ['nav-app-01',      'Servidor Aplicacion Dynamics NAV (primario)',       'Por confirmar', 'NST NAV — sesiones de usuarios + servicios web'],
        ['nav-app-02',      'Servidor Aplicacion Dynamics NAV (secundario)',     'Por confirmar', 'NST NAV — balanceo de carga / alta disponibilidad'],
        ['adds-replica',    'Controlador de Dominio AD DS — Replica Azure',      'Por confirmar', 'DNS + autenticacion Kerberos en la zona Azure'],
    ]
    for row in server_rows:
        r = add_table_row(servers_tbl, row)
        # Destacar los servidores nuevos (pendientes de inventario)
        if 'Por confirmar' in row[2]:
            for cell in r.cells:
                set_cell_bg(cell, 'FFF3CD')

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('NOTA: ')
    r.bold = True
    p.add_run(
        'Los servidores nav-app-01, nav-app-02 y adds-replica estan pendientes de inventario '
        '(RAM, CPU, version de SO y NAV). El sizing de estas VMs se realizara en una segunda fase '
        'una vez completados los ciclos de monitorización SQL.'
    )

    add_heading(doc, '2.2 Arquitectura de red en Azure (IaaS)', 2)
    doc.add_paragraph(
        'Al migrar toda la pila de aplicacion a Azure IaaS, la latencia entre los servidores NAV '
        'y SQL Server sera minima (red interna de Azure, < 1 ms dentro de la misma region y VNet). '
        'El dominio AD DS se extiende a Azure mediante una replica del controlador de dominio, '
        'evitando dependencias de autenticacion criticas con el entorno on-premises durante la '
        'operacion normal en Azure.'
    )

    arch_bullets = [
        'Azure Virtual Network (VNet) con subredes segregadas: aplicacion, datos, administracion.',
        'SQL Server en VMs con discos Premium SSD o Ultra Disk segun los IOPS definitivos del sizing.',
        'Servidores NAV en VMs Windows Server 2022 con el mismo NST y version de Dynamics NAV actual.',
        'AD DS replica en VM ligera (Standard_D2s_v5 orientativo) — solo lectura/DNS desde Azure.',
        'Conectividad hibrida con on-premises via Azure VPN Gateway o ExpressRoute durante la transicion.',
    ]
    for b in arch_bullets:
        add_bullet(doc, b)

    # -----------------------------------------------------------------------
    # 3. ANALISIS PREVIOS (13/04/2026)
    # -----------------------------------------------------------------------
    add_heading(doc, '3. Analisis Realizados el 13/04/2026', 1)
    doc.add_paragraph(
        'Con caracter previo al analisis de workload, se realizaron tres auditorias tecnicas que '
        'identificaron incidencias criticas de operacion y seguridad.'
    )

    for i, report in enumerate(PREV_REPORTS, 1):
        add_heading(doc, f'3.{i} {report["title"]}', 2)

        risk_color = {'CRITICO': 'FF4444', 'ALTO': 'FF8C00', 'MEDIO': 'FFD700'}.get(report['risk'], 'CCCCCC')
        p = doc.add_paragraph()
        r = p.add_run(f'Riesgo: {report["risk"]}  ')
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(risk_color)

        for finding in report['findings']:
            add_bullet(doc, finding)

    # -----------------------------------------------------------------------
    # 4. CICLO 1 DE MONITORIZACION (13-14/04/2026)
    # -----------------------------------------------------------------------
    add_heading(doc, '4. Ciclo 1 de Monitorizacion — Resultados (13-14/04/2026)', 1)

    if not monitoring_data:
        p = doc.add_paragraph(
            'NOTA: No se encontraron datos de monitorización en docs/data/. '
            'Ejecuta primero: .\\scripts\\Collect-MonitoringData.ps1'
        )
        for run in p.runs:
            run.italic = True

        # Incluir datos conocidos del checkpoint ~9:00 AM del 14/04
        doc.add_paragraph(
            'Los siguientes datos corresponden al checkpoint recogido el 14/04/2026 a las ~09:00 '
            '(60% del ciclo completado, ~855-864 minutos de 1440):'
        )

        chk_tbl = doc.add_table(rows=1, cols=6)
        chk_tbl.style = 'Table Grid'
        chk_hdrs = ['Servidor', 'Samples', 'Elapsed (min)', 'Top Wait', 'Usuarios activos', 'Buffer Pool']
        for i, h in enumerate(chk_hdrs):
            chk_tbl.rows[0].cells[i].text = h
            chk_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(chk_tbl.rows[0].cells[i], 'D6E4F0')

        chk_data = [
            ['srvplenoilfs',   '333', '855', 'SOS_WORK_DISPATCHER', '28', '131 / 240 GB'],
            ['sql-restore-001','314', '864', 'SOS_WORK_DISPATCHER', '56', '131 / 240 GB'],
        ]
        for row in chk_data:
            add_table_row(chk_tbl, row)

    else:
        doc.add_paragraph(
            'Los siguientes datos corresponden al Ciclo 1 completo de 24 horas de monitorización '
            '(13/04/2026 18:18 — 14/04/2026 18:24).'
        )

        for server_key, srv_data in monitoring_data.items():
            add_heading(doc, f'4.x Servidor: {server_key}', 2)
            st = srv_data['stats']
            if not st:
                doc.add_paragraph(f'No se pudieron calcular estadísticas para {server_key}.')
                continue

            # Tabla de metricas
            metrics_tbl = doc.add_table(rows=1, cols=3)
            metrics_tbl.style = 'Table Grid'
            for i, h in enumerate(['Metrica', 'Promedio', 'Maximo']):
                metrics_tbl.rows[0].cells[i].text = h
                metrics_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                set_cell_bg(metrics_tbl.rows[0].cells[i], 'D6E4F0')

            metrics_rows = [
                ['CPU SQL Server (%)',      f"{st['cpu_sql_avg']}%",       f"{st['cpu_sql_max']}%"],
                ['CPU Total (%)',           f"{st['cpu_total_avg']}%",     f"{st['cpu_total_max']}%"],
                ['Buffer Pool (GB)',        f"{st['buf_pool_avg_gb']} GB", f"{st['buf_pool_max_gb']} GB"],
                ['Memoria total SQL (GB)',  f"—",                          f"{st['total_mem_max_gb']} GB"],
                ['Conexiones activas',      f"{st['conn_avg']}",           f"{st['conn_max']}"],
                ['Requests activos',        f"{st['req_avg']}",            f"{st['req_max']}"],
                ['Procesos bloqueados',     f"—",                          f"{st['blocked_max']} (eventos: {st['blocked_events']})"],
                ['Memory grants pendientes','—',                           f"{st['mem_grants_max']}"],
            ]
            for row in metrics_rows:
                add_table_row(metrics_tbl, row)

            doc.add_paragraph()

            # Top waits
            add_heading(doc, f'Top Wait Types — {server_key}', 3)
            waits_tbl = doc.add_table(rows=1, cols=3)
            waits_tbl.style = 'Table Grid'
            for i, h in enumerate(['Wait Type', 'Tiempo acumulado (ms)', 'Interpretacion']):
                waits_tbl.rows[0].cells[i].text = h
                waits_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                set_cell_bg(waits_tbl.rows[0].cells[i], 'D6E4F0')

            wait_interp = {
                'SOS_WORK_DISPATCHER':  'Normal — procesamiento interno del scheduler',
                'CXPACKET':             'Paralelismo — posible desajuste de MaxDOP',
                'PAGEIOLATCH_SH':       'I/O de lectura — fragmentacion o disco lento',
                'PAGEIOLATCH_EX':       'I/O de escritura — checkpoint o rebuild',
                'LCK_M_S':              'Lock de lectura — contención de concurrencia',
                'LCK_M_X':              'Lock exclusivo — escrituras concurrentes',
                'WRITELOG':             'Log de transacciones — I/O de log',
                'ASYNC_NETWORK_IO':     'Red — cliente lento procesando resultados',
                'RESOURCE_SEMAPHORE':   'Memory grant pendiente — presion de memoria',
                'PREEMPTIVE_OS_FILEOPS':'Operacion de archivo del SO — backup/I/O',
            }
            for wt, wms in st['top_waits']:
                interp = wait_interp.get(wt, 'Ver documentacion de Microsoft para este wait type')
                add_table_row(waits_tbl, [wt, f"{round(wms/1000, 1)} s", interp])

    # -----------------------------------------------------------------------
    # 5. JOBS SQL AGENT CRITICOS
    # -----------------------------------------------------------------------
    add_heading(doc, '5. Jobs SQL Agent Identificados', 1)
    doc.add_paragraph(
        'Los siguientes jobs son criticos para el sizing porque determinan la carga pico en cada servidor. '
        'Los marcados como NO capturados requieren monitorización adicional para obtener metricas reales.'
    )

    jobs_tbl = doc.add_table(rows=1, cols=6)
    jobs_tbl.style = 'Table Grid'
    jobs_hdrs = ['Job', 'Servidor', 'Horario', 'Duracion', 'Prioridad', 'Ciclo 1']
    for i, h in enumerate(jobs_hdrs):
        jobs_tbl.rows[0].cells[i].text = h
        jobs_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(jobs_tbl.rows[0].cells[i], 'D6E4F0')

    for job in SQL_JOBS:
        captured_str = '✓ Sí' if job['captured'] else '✗ No'
        duration_str = f"{job['duration_h']:.1f}h"
        row = jobs_tbl.add_row()
        vals = [job['name'], job['server'], job['schedule'], duration_str, job['priority'], captured_str]
        for i, val in enumerate(vals):
            row.cells[i].text = val
            if not job['captured'] and i in (0, 5):
                row.cells[i].paragraphs[0].runs[0].bold = True
        if not job['captured']:
            for cell in row.cells:
                set_cell_bg(cell, 'FFF3CD')  # Amarillo suave

    # -----------------------------------------------------------------------
    # 6. EVALUACION DE REPRESENTATIVIDAD
    # -----------------------------------------------------------------------
    add_heading(doc, '6. Evaluacion de Representatividad — ¿Son Suficientes 24 Horas?', 1)

    p = doc.add_paragraph()
    r = p.add_run('Conclusion: Las 24 horas del Ciclo 1 NO son suficientemente representativas ')
    r.bold = True
    p.add_run(
        'del workload real de produccion. Se identifican tres gaps significativos que pueden '
        'impactar materialmente el sizing de las instancias Azure:'
    )

    gaps = [
        ('IndicesEstadisticas (sql-restore-001)',
         '25h 43min de duracion, ejecuta Dom/Mar/Mie/Vie a las 18:30. '
         'El Ciclo 1 capturó solo el inicio de este job. Es el proceso mas largo del entorno '
         'y probablemente define el dimensionamiento de almacenamiento e IOPS.'),
        ('ProcesarCubosdomingo (srvplenoilfs)',
         '6h 21min de duracion, ejecuta sabados y lunes a las 01:30. '
         'Requiere monitorización de fin de semana para capturar la carga real.'),
        ('Pico de fin de mes',
         'Los usuarios preparan informes de cierre mensual los ultimos dias del mes (aprox. 27-30 de cada mes). '
         'Este pico no es capturable en una semana tipica y requiere monitorización especifica a finales de abril.'),
    ]

    for title, desc in gaps:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{title}: ')
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Factor de seguridad recomendado para el sizing previo a monitorización completa: ').bold = True
    p.add_run('x1,3 sobre los valores pico del Ciclo 1. Este factor debe revisarse una vez completados '
              'los ciclos de fin de semana y fin de mes.')

    # -----------------------------------------------------------------------
    # 7. PLAN DE MONITORIZACION EXTENDIDA
    # -----------------------------------------------------------------------
    add_heading(doc, '7. Plan de Monitorizacion Extendida', 1)
    doc.add_paragraph(
        'Para obtener una fotografia completa del workload real, se propone el siguiente plan '
        'de monitorización adicional antes de proceder al sizing definitivo:'
    )

    cycles_tbl = doc.add_table(rows=1, cols=5)
    cycles_tbl.style = 'Table Grid'
    cycles_hdrs = ['Ciclo', 'Servidor(es)', 'Inicio', 'Duracion', 'Objetivo']
    for i, h in enumerate(cycles_hdrs):
        cycles_tbl.rows[0].cells[i].text = h
        cycles_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(cycles_tbl.rows[0].cells[i], 'D6E4F0')

    today_str = date.today().strftime('%d/%m/%Y')
    friday_str = 'Viernes 22:00'
    cycle_rows = [
        ['Ciclo 1 ✓', 'Ambos', f'13/04/2026 18:18', '24h (completado)', 'Baseline diario — jobs nocturnos'],
        ['Ciclo 2',   'Ambos', f'{today_str} 18:35', '48h',
         'Captura IndicesEstadisticas completo (>25h desde martes 18:30)'],
        ['Ciclo 3',   'srvplenoilfs', f'{friday_str}', '48h',
         'Captura ProcesarCubosdomingo (sabado 01:30) y carga de fin de semana'],
        ['Ciclo 4',   'Ambos', '27/04/2026 08:00', '72-96h',
         'Captura pico de fin de mes — usuarios con informes de cierre'],
    ]
    for row in cycle_rows:
        r = add_table_row(cycles_tbl, row)
        if '✓' in row[0]:
            set_cell_bg(r.cells[0], 'D4EDDA')

    doc.add_paragraph()
    add_heading(doc, '7.1 Automatizacion', 2)
    doc.add_paragraph(
        'Los Ciclos 2 y 3 se han programado como Scheduled Tasks directamente en los servidores '
        'remotos de produccion. Las tareas se ejecutaran en los propios servidores '
        'independientemente del estado del equipo del administrador. '
        'El Ciclo 4 debera programarse manualmente a finales de abril.'
    )

    # -----------------------------------------------------------------------
    # 8. CONSIDERACIONES PARA EL SIZING AZURE IAAS
    # -----------------------------------------------------------------------
    add_heading(doc, '8. Consideraciones para el Sizing Azure IaaS', 1)

    add_heading(doc, '8.1 Modelo de despliegue: IaaS (lift & shift)', 2)
    doc.add_paragraph(
        'La migracion adopta el modelo IaaS: cada servidor on-premises se convierte en una Azure VM '
        'equivalente, manteniendo el mismo sistema operativo, version de SQL Server y configuracion '
        'de Dynamics NAV. Esto garantiza compatibilidad total con los procedimientos almacenados '
        'y los procesos automatizados actuales, sin necesidad de adaptar codigo.'
    )
    iaas_bullets = [
        'SQL Server on Azure VM: version identica al servidor actual, licencia via Azure Hybrid Benefit (AHUB) '
        'para reducir coste si dispone de licencias SA activas.',
        'Dynamics NAV NST: misma version instalada en VMs Windows Server equivalentes — sin recompilacion.',
        'AD DS replica: extiende el dominio existente a Azure, Kerberos funciona sin cambios para NAV y SQL.',
        'No se requiere reescritura de stored procedures, jobs SQL Agent ni logica de integracion.',
    ]
    for b in iaas_bullets:
        add_bullet(doc, b)

    add_heading(doc, '8.2 Sizing preliminar SQL Server (a revisar tras Ciclos 2-4)', 2)
    doc.add_paragraph(
        'Con los datos del Ciclo 1 y aplicando el factor de seguridad x1,3, los rangos preliminares '
        'para las VMs SQL son:'
    )

    sizing_tbl = doc.add_table(rows=1, cols=5)
    sizing_tbl.style = 'Table Grid'
    sizing_hdrs = ['Servidor', 'RAM medida', 'RAM Azure (+30%)', 'CPU medido', 'SKU orientativo Azure VM']
    for i, h in enumerate(sizing_hdrs):
        sizing_tbl.rows[0].cells[i].text = h
        sizing_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(sizing_tbl.rows[0].cells[i], 'D6E4F0')

    sizing_rows = [
        ['srvplenoilfs',   '131 GB usados / 240 GB totales', '~170 GB', 'Ver Ciclo 2+', 'Standard_E48ds_v5 (384 GB) o Mv2 series'],
        ['sql-restore-001','131 GB usados / 240 GB totales', '~170 GB', 'Ver Ciclo 2+', 'Standard_E48ds_v5 (384 GB) o Mv2 series'],
    ]
    for row in sizing_rows:
        add_table_row(sizing_tbl, row)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IMPORTANTE: ').bold = True
    p.add_run(
        'El sizing definitivo debe realizarse tras completar los Ciclos 2, 3 y 4 de monitorización. '
        'Los valores anteriores son orientativos y pueden variar significativamente al capturar '
        'IndicesEstadisticas (job de 25h) y los picos de fin de mes.'
    )

    add_heading(doc, '8.3 Sizing preliminar Capa de Aplicacion NAV', 2)
    doc.add_paragraph(
        'Los servidores de aplicacion NAV estan pendientes de inventario completo. '
        'Como referencia orientativa para entornos con carga media-alta en Dynamics NAV:'
    )
    nav_sizing_tbl = doc.add_table(rows=1, cols=4)
    nav_sizing_tbl.style = 'Table Grid'
    for i, h in enumerate(['Servidor', 'Rol', 'SKU orientativo Azure VM', 'Notas']):
        nav_sizing_tbl.rows[0].cells[i].text = h
        nav_sizing_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(nav_sizing_tbl.rows[0].cells[i], 'D6E4F0')

    nav_rows = [
        ['nav-app-01', 'NST primario',    'Standard_D8s_v5 (32 GB) — revisar con inventario', 'Mismo OS/NAV que on-prem'],
        ['nav-app-02', 'NST secundario',  'Standard_D8s_v5 (32 GB) — revisar con inventario', 'Balanceo / HA'],
        ['adds-replica','AD DS replica',  'Standard_D2s_v5 (8 GB) — suficiente para replica',  'DNS en VNet + Kerberos'],
    ]
    for row in nav_rows:
        r = add_table_row(nav_sizing_tbl, row)
        for cell in r.cells:
            set_cell_bg(cell, 'FFF3CD')  # Pendientes de confirmar

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Accion pendiente: ').bold = True
    p.add_run(
        'Recopilar el inventario de nav-app-01, nav-app-02 y adds-replica '
        '(CPU, RAM, OS, version NAV, numero de usuarios concurrentes NST) para '
        'ajustar el sizing de estas VMs antes de presentar la propuesta economica.'
    )

    add_heading(doc, '8.4 Nota especial — Fragmentacion y IOPS', 2)
    doc.add_paragraph(
        'La fragmentacion critica identificada en srvplenoilfs (98,46% en PK_HechosVentasPOS) '
        'genera un consumo de IOPS artificialmente elevado. '
        'Se recomienda ejecutar el REBUILD INDEX ANTES de realizar las mediciones definitivas '
        'de IOPS, para no sobredimensionar el almacenamiento Azure SQL en base a un estado '
        'patologico del servidor. En Azure IaaS, esto afecta directamente al tipo de disco '
        '(Premium SSD P-series vs Ultra Disk) y su impacto economico es significativo.'
    )

    add_heading(doc, '8.5 Azure Hybrid Benefit — Optimizacion de Costes', 2)
    doc.add_paragraph(
        'Si Plenergy dispone de licencias SQL Server y Windows Server con Software Assurance (SA) activo, '
        'puede aplicar Azure Hybrid Benefit para reducir hasta un 40-85% el coste de las VMs SQL. '
        'Se recomienda verificar el estado de las licencias antes de la propuesta economica final.'
    )

    # -----------------------------------------------------------------------
    # 9. PROXIMOS PASOS
    # -----------------------------------------------------------------------
    add_heading(doc, '9. Proximos Pasos', 1)

    steps = [
        ('Inmediato (esta semana)',
         'Recopilar inventario de nav-app-01, nav-app-02 y adds-replica: CPU, RAM, OS, version NAV, '
         'numero de usuarios concurrentes NST. Necesario para completar el sizing de la propuesta.'),
        ('Inmediato (esta semana)',
         'Ejecutar REBUILD INDEX ONLINE sobre PK_HechosVentasPOS en ventana de mantenimiento '
         '(noche del martes/miercoles). Esto requiere SQL Server Enterprise Edition para ONLINE=ON.'),
        ('Inmediato (esta semana)',
         'Deshabilitar login "sa" y revisar los 16 logins con sysadmin. Documentar cambios en '
         'sistema de control de cambios.'),
        ('Martes 14/04 — Automatico',
         'Inicio del Ciclo 2 de monitorización (48h) en ambos servidores SQL. '
         'La tarea esta programada en los servidores remotos.'),
        ('Viernes 18/04 — Automatico',
         'Inicio del Ciclo 3 de monitorización en srvplenoilfs (captura fin de semana, '
         'job ProcesarCubosdomingo).'),
        ('Semana del 22/04',
         'Analizar resultados Ciclos 2 y 3. Presentar sizing revisado SQL + sizing NAV/AD DS '
         'con inventario completado.'),
        ('Ultima semana de abril (27-30/04)',
         'Ciclo 4 de monitorización para capturar pico de fin de mes. '
         'Programar manualmente a finales de abril.'),
        ('Mayo 2026',
         'Sizing definitivo completo (5 VMs). Propuesta economica con Azure Hybrid Benefit. '
         'Inicio del plan de migracion IaaS en fases.'),
    ]

    steps_tbl = doc.add_table(rows=1, cols=3)
    steps_tbl.style = 'Table Grid'
    for i, h in enumerate(['Plazo', 'Accion', 'Responsable']):
        steps_tbl.rows[0].cells[i].text = h
        steps_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(steps_tbl.rows[0].cells[i], 'D6E4F0')

    for plazo, accion in steps:
        add_table_row(steps_tbl, [plazo, accion, 'Equipo de infraestructura'])

    # -----------------------------------------------------------------------
    # FOOTER con metadata del informe
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f'Informe generado automaticamente el {datetime.now().strftime("%d/%m/%Y %H:%M")} | '
        f'Datos: {len(monitoring_data)} servidor(es) analizados'
    )
    meta_run.font.size  = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # -----------------------------------------------------------------------
    # Guardar
    # -----------------------------------------------------------------------
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\n[OK] Informe guardado: {output_path}")
    print(f"     Tamano: {output_path.stat().st_size / 1024:.0f} KB")


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description='Genera el informe Word de migracion Azure para Plenergy.'
    )
    parser.add_argument(
        '--output', '-o',
        default=str(SCRIPT_DIR / f'Informe-Planificacion-Migracion-Azure-Plenergy-{date.today().strftime("%Y%m%d")}.docx'),
        help='Ruta del archivo Word de salida (default: docs/Informe-...docx)'
    )
    parser.add_argument(
        '--data-dir', '-d',
        default=str(DEFAULT_DATA_DIR),
        help='Directorio con los JSON de monitorización (default: docs/data/)'
    )
    parser.add_argument(
        '--cycle', '-c',
        choices=['checkpoint', 'cycle1', 'cycle2', 'weekend'],
        default='checkpoint',
        help='Ciclo de datos a incluir (default: checkpoint)'
    )
    args = parser.parse_args()

    print("=" * 60)
    print("  Generador de Informe — Migracion Azure Plenergy")
    print("=" * 60)

    data_dir = Path(args.data_dir)
    print(f"\nBuscando datos en: {data_dir}")
    monitoring_data = load_monitoring_data(data_dir)

    if not monitoring_data:
        print("\nWARN: No se encontraron datos de monitorización.")
        print(f"  Ejecuta primero: .\\scripts\\Collect-MonitoringData.ps1")
        print(f"  El informe se generara con datos del checkpoint conocido (09:00 AM 14/04).\n")

    output_path = Path(args.output)
    print(f"\nGenerando informe Word: {output_path}")
    generate_report(monitoring_data, output_path)


if __name__ == '__main__':
    main()
